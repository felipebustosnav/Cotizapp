import docx
import sys
import io

# Set UTF-8 encoding for stdout
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def extract_full_docx(docx_path):
    doc = docx.Document(docx_path)
    
    full_text = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            if any(row_text):
                full_text.append(" | ".join(row_text))
    
    return "\n".join(full_text)

if __name__ == "__main__":
    docx_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    text = extract_full_docx(docx_path)
    
    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        print(f"Texto extraído y guardado en {output_path}")
    else:
        print(text)
